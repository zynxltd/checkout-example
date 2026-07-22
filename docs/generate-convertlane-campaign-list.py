#!/usr/bin/env python3
"""Generate ConvertLane UK Campaign Portfolio docx (1 Click Wonder–style rate card)."""

import importlib.util
import math
import re
import shutil
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "convertlane-campaign-list-assets"
OUT_DOCX = ROOT / "ConvertLane-CampaignPortfolio-June2026.docx"
OUT_DOWNLOADS = Path.home() / "Downloads" / "ConvertLane-CampaignPortfolio-June2026.docx"
REF_DOCX = Path.home() / "Downloads" / "GeneralCampaignListJune2026.docx"

DPI = 300
PAGE_W = int(8.27 * DPI)
PAGE_H = int(11.69 * DPI)

NAVY = (10, 18, 32)
DEEP = (6, 12, 24)
TEAL = (0, 180, 170)
GOLD = (212, 175, 95)
GOLD_LIGHT = (235, 210, 150)
WHITE = (255, 255, 255)
MUTED = (140, 155, 175)

WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def load_offers():
    spec = importlib.util.spec_from_file_location(
        "offers", ROOT / "convertlane-uk-finance-offers.py"
    )
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    uk = [r for r in mod.ROWS if r[2] == "UK"]
    groups = {}
    for r in uk:
        groups.setdefault(r[3], []).append(r)
    order = [
        "Finance",
        "Loans",
        "iGaming",
        "Betting",
        "Energy",
        "Sweepstakes",
        "Surveys",
        "Other",
    ]
    return [(v, groups[v]) for v in order if v in groups]


def simplify_reward(model: str) -> str:
    m = model.upper().replace(" ", "")
    for key in ("REVSHARE", "CPS", "SOI", "CPL", "CPA"):
        if key in m:
            return {"REVSHARE": "RevShare"}.get(key, key)
    return model.split("/")[0].strip()


def format_commission(payout: str, model: str) -> str:
    p = payout.strip()
    if not p or p.upper() == "TBC":
        return "Tiered" if "revshare" in model.lower() else "TBC"
    if "%" in p or "RevShare" in p:
        return p.replace("$", "").replace("€", "")
    if p.startswith("~"):
        p = p[1:]
    p = p.replace("$", "£").replace("€", "£")
    if "£" not in p and p and p[0].isdigit():
        return f"£{p}"
    return p


def row_to_campaign(row):
    offer, sub = row[5], row[4]
    name = offer if sub in ("", offer) else f"{offer} — {sub}"
    return {
        "campaign": name[:85],
        "commission": format_commission(row[8], row[7]),
        "reward": simplify_reward(row[7]),
        "market": row[2],
        "preview": "Landing Page",
    }


# ── 300 DPI artwork ──────────────────────────────────────────────────────────
def _font(size, bold=False):
    paths = [
        ("/System/Library/Fonts/Supplemental/Avenir Next.ttc", 1 if bold else 0),
        ("/System/Library/Fonts/Helvetica.ttc", 1 if bold else 0),
        ("/Library/Fonts/Arial Bold.ttf", 0),
        ("/Library/Fonts/Arial.ttf", 0),
    ]
    for path, idx in paths:
        if Path(path).exists():
            try:
                return ImageFont.truetype(path, size, index=idx)
            except Exception:
                try:
                    return ImageFont.truetype(path, size)
                except Exception:
                    pass
    return ImageFont.load_default()


def _draw_logo(draw, cx, cy, radius, color=GOLD):
    for i, r in enumerate([radius, radius * 0.72, radius * 0.44]):
        c = tuple(max(0, int(color[j] * (1 - i * 0.12))) for j in range(3))
        draw.arc(
            [cx - r, cy - r, cx + r, cy + r],
            start=200,
            end=340,
            fill=c,
            width=max(5, int(radius * 0.09)),
        )
    for angle in (220, 260, 300):
        rad = math.radians(angle)
        x = cx + radius * 0.85 * math.cos(rad)
        y = cy + radius * 0.85 * math.sin(rad)
        draw.ellipse([x - 16, y - 16, x + 16, y + 16], fill=color)


def make_cover(path: Path):
    img = Image.new("RGB", (PAGE_W, PAGE_H), NAVY)
    draw = ImageDraw.Draw(img)
    for y in range(PAGE_H):
        t = y / PAGE_H
        col = (
            int(8 + 14 * (1 - t)),
            int(22 + 10 * (1 - t)),
            int(42 + 8 * (1 - t)),
        )
        draw.line([(0, y), (PAGE_W, y)], fill=col)

    # teal glow orbs
    glow = Image.new("RGBA", (PAGE_W, PAGE_H), (0, 0, 0, 0))
    gd = ImageDraw.Draw(glow)
    for cx, cy, rad in [(PAGE_W * 0.78, PAGE_H * 0.22, 700), (PAGE_W * 0.18, PAGE_H * 0.65, 550)]:
        for r in range(rad, 0, -10):
            gd.ellipse([cx - r, cy - r, cx + r, cy + r], fill=(0, 160, 150, int(35 * r / rad)))
    img = Image.alpha_composite(img.convert("RGBA"), glow).convert("RGB")
    draw = ImageDraw.Draw(img)

    # abstract performance chart
    bx, by = int(PAGE_W * 0.1), int(PAGE_H * 0.55)
    for i, bh in enumerate([140, 240, 180, 320, 260, 380]):
        draw.rectangle([bx + i * 58, by - bh, bx + i * 58 + 44, by], fill=(0, 130, 125))
    draw.text((bx - 10, by + 20), "£", fill=GOLD, font=_font(130, bold=True))

    # routing lines
    for i in range(14):
        draw.line(
            [
                (PAGE_W * 0.55, PAGE_H * 0.3 + i * 18),
                (PAGE_W * 0.92, PAGE_H * 0.38 + i * 22),
            ],
            fill=(0, 100, 95),
            width=2,
        )

    lx = PAGE_W // 2
    _draw_logo(draw, lx, int(PAGE_H * 0.11), 100)
    bf = _font(80, bold=True)
    brand = "ConvertLane"
    draw.text((lx - draw.textlength(brand, font=bf) / 2, PAGE_H * 0.11 + 120), brand, fill=GOLD_LIGHT, font=bf)
    sf = _font(34)
    tag = "Performance Affiliate Network"
    draw.text((lx - draw.textlength(tag, font=sf) / 2, PAGE_H * 0.11 + 215), tag, fill=MUTED, font=sf)

    tf = _font(100, bold=True)
    t1 = "CAMPAIGN PORTFOLIO"
    draw.text((lx - draw.textlength(t1, font=tf) / 2, PAGE_H * 0.76), t1, fill=WHITE, font=tf)
    df = _font(68, bold=True)
    t2 = "JUNE 2026"
    draw.text((lx - draw.textlength(t2, font=df) / 2, PAGE_H * 0.76 + 120), t2, fill=GOLD, font=df)

    ff = _font(30)
    ft = "UK · Finance · Loans · iGaming · Betting · Energy"
    draw.text((lx - draw.textlength(ft, font=ff) / 2, PAGE_H * 0.91), ft, fill=MUTED, font=ff)

    img.save(path, "JPEG", quality=98, dpi=(DPI, DPI), subsampling=0)


def make_inner_bg(path: Path):
    img = Image.new("RGB", (PAGE_W, PAGE_H), (4, 6, 10))
    draw = ImageDraw.Draw(img)
    for cx, cy in [(PAGE_W * 0.98, -PAGE_H * 0.08), (-PAGE_W * 0.08, PAGE_H * 1.06)]:
        for r in range(350, 2000, 130):
            draw.arc([cx - r, cy - r, cx + r, cy + r], 0, 360, fill=(28, 32, 42), width=4)
    _draw_logo(draw, 130, 110, 60)
    draw.text((210, 60), "ConvertLane", fill=GOLD_LIGHT, font=_font(52, bold=True))
    img.save(path, "JPEG", quality=98, dpi=(DPI, DPI), subsampling=0)


# ── docx styling (dark rate-card theme) ──────────────────────────────────────
def set_cell_shading(cell, hex_color: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shd)


def set_cell_text(cell, text: str, *, bold=False, color="FFFFFF", size=9, underline=False, align="left"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if underline:
        run.underline = True
    rgb = RGBColor(int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))
    run.font.color.rgb = rgb


def add_bg_to_header(section, image_path: Path):
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.clear()
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(8.27), height=Inches(11.69))


def add_table(doc, section_name: str, campaigns: list):
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    headers = ["Campaign", "Commission", "Reward", "Market", "Preview"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, color="FFFFFF", size=10, align="center")
        set_cell_shading(table.rows[0].cells[i], "000000")

    sec = table.add_row()
    set_cell_text(sec.cells[0], section_name, bold=True, color="D4AF5F", size=11, align="center")
    for i in range(1, 5):
        set_cell_text(sec.cells[i], "", color="FFFFFF")
        set_cell_shading(sec.cells[i], "1A1A1A")
    set_cell_shading(sec.cells[0], "1A1A1A")

    for c in campaigns:
        row = table.add_row()
        vals = [c["campaign"], c["commission"], c["reward"], c["market"], c["preview"]]
        for i, v in enumerate(vals):
            if i == 0:
                set_cell_text(row.cells[i], v, color="0563C1", size=9, underline=True)
            elif i == 4:
                set_cell_text(row.cells[i], v, color="0563C1", size=9, underline=True, align="center")
            else:
                set_cell_text(row.cells[i], v, color="FFFFFF", size=9, align="center" if i > 0 else "left")
            set_cell_shading(row.cells[i], "000000")

    doc.add_paragraph("")


def add_contact_block(doc):
    doc.add_page_break()
    blocks = [
        ("PUBLISHER LOGIN", True, "D4AF5F", 14),
        ("https://convertlane.co.uk/publishers", False, "0563C1", 11),
        ("PUBLISHER SIGN UP", True, "D4AF5F", 14),
        ("https://convertlane.co.uk/publishers", False, "0563C1", 11),
        ("", False, "FFFFFF", 6),
        ("partners@convertlane.co.uk", False, "FFFFFF", 11),
        ("https://convertlane.co.uk", False, "0563C1", 11),
        ("", False, "FFFFFF", 6),
        (
            "ConvertLane Ltd · Registered in England & Wales\n"
            "Net-30 payouts on the 15th · Min £100 · Vetted partners only",
            False,
            "888888",
            9,
        ),
    ]
    for text, bold, color, size in blocks:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if not text:
            continue
        run = p.add_run(text + "\n")
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))
        if color == "0563C1":
            run.underline = True


def build_from_template(groups, cover_jpg, inner_jpg):
    """Clone reference docx shell; swap 300 DPI backgrounds + table text + branding."""
    if not REF_DOCX.exists():
        return False

    tmp = ASSETS / "_tpl"
    if tmp.exists():
        shutil.rmtree(tmp)
    tmp.mkdir()
    with zipfile.ZipFile(REF_DOCX, "r") as z:
        z.extractall(tmp)

    shutil.copy(cover_jpg, tmp / "word/media/image4.jpg")
    shutil.copy(inner_jpg, tmp / "word/media/image5.jpg")

    # flatten campaigns in vertical order
    all_rows = []
    for vertical, rows in groups:
        all_rows.append(("__SECTION__", vertical))
        for r in rows:
            c = row_to_campaign(r)
            all_rows.append(
                ("__DATA__", c["campaign"], c["commission"], c["reward"], c["market"], c["preview"])
            )

    xml_path = tmp / "word/document.xml"
    xml = xml_path.read_text(encoding="utf-8")

    # branding replacements
    reps = {
        "1 Click Wonder": "ConvertLane",
        "OneClickWonder": "ConvertLane",
        "www.1clickwonder.com": "https://convertlane.co.uk",
        "teams@1clickwonder.com": "partners@convertlane.co.uk",
        "1 Click Wonder is a trading style of Ventures London LtdRegistered in England No: 11132256": (
            "ConvertLane Ltd · Registered in England & Wales"
        ),
        "1 Click Wonder is a trading style of Ventures London Ltd": "ConvertLane Ltd",
        "Registered in England No: 11132256": "",
    }
    for old, new in reps.items():
        xml = xml.replace(old, new)

    # replace table cell texts in document order
    data_iter = iter(all_rows)

    def repl_section(match):
        try:
            item = next(data_iter)
        except StopIteration:
            return match.group(0)
        if item[0] != "__SECTION__":
            return match.group(0)
        return f"<w:t>{item[1]}</w:t>"

    def repl_data_row(cells_texts):
        """Replace 5 consecutive w:t in a row — handled via sequential replace."""
        pass

    # Sequential replacement of campaign row texts (skip header rows)
    t_pattern = re.compile(
        r"(<w:tr[^>]*>[\s\S]*?</w:tr>)",
    )
    tables = re.findall(r"<w:tbl>[\s\S]*?</w:tbl>", xml)
    new_tables = []
    row_q = list(all_rows)

    for tbl in tables:
        rows = re.findall(r"<w:tr[\s\S]*?</w:tr>", tbl)
        if not rows or "Campaign" not in rows[0]:
            new_tables.append(tbl)
            continue
        out_rows = [rows[0]]  # keep header
        ri = 1
        while ri < len(rows) and row_q:
            item = row_q[0]
            if item[0] == "__SECTION__":
                row_q.pop(0)
                # reuse section row template if available else data row
                tmpl = rows[ri] if ri < len(rows) else rows[min(1, len(rows) - 1)]
                new_row = re.sub(
                    r"<w:t[^>]*>[^<]*</w:t>",
                    lambda m, _v=item[1]: f"<w:t>{_v}</w:t>",
                    tmpl,
                    count=1,
                )
                # blank remaining cells in section row
                rest = new_row.split("</w:t>", 1)
                if len(rest) == 2:
                    new_row = rest[0] + "</w:t>" + re.sub(
                        r"<w:t[^>]*>[^<]*</w:t>", "<w:t></w:t>", rest[1]
                    )
                out_rows.append(new_row)
                ri += 1
            elif item[0] == "__DATA__":
                row_q.pop(0)
                if ri >= len(rows):
                    break
                tmpl = rows[ri]
                values = list(item[1:])
                vi = 0

                def sub_t(m):
                    nonlocal vi
                    if vi >= len(values):
                        return m.group(0)
                    val = values[vi]
                    vi += 1
                    return f"<w:t>{val}</w:t>"

                new_row = re.sub(r"<w:t[^>]*>[^<]*</w:t>", sub_t, tmpl)
                out_rows.append(new_row)
                ri += 1
            else:
                ri += 1
        new_tbl = "<w:tbl>" + "".join(
            re.search(r"(<w:tblPr>[\s\S]*?</w:tblPr>)", tbl).group(1)
            if re.search(r"(<w:tblPr>[\s\S]*?</w:tblPr>)", tbl)
            else ""
        )
        new_tbl += "".join(out_rows) + "</w:tbl>"
        new_tables.append(new_tbl)

    it = iter(new_tables)
    xml = re.sub(r"<w:tbl>[\s\S]*?</w:tbl>", lambda _: next(it), xml)

    xml_path.write_text(xml, encoding="utf-8")

    # update core metadata
    core = tmp / "docProps/core.xml"
    if core.exists():
        c = core.read_text(encoding="utf-8")
        c = re.sub(r"<dc:title>[^<]*</dc:title>", "<dc:title>ConvertLane Campaign Portfolio June 2026</dc:title>", c)
        c = re.sub(r"<dc:creator>[^<]*</dc:creator>", "<dc:creator>ConvertLane</dc:creator>", c)
        core.write_text(c, encoding="utf-8")

    # repack
    if OUT_DOCX.exists():
        OUT_DOCX.unlink()
    with zipfile.ZipFile(OUT_DOCX, "w", zipfile.ZIP_DEFLATED) as zout:
        for fp in sorted(tmp.rglob("*")):
            if fp.is_file():
                zout.write(fp, fp.relative_to(tmp).as_posix())

    shutil.rmtree(tmp)
    return True


def build_fresh_docx(groups, cover_jpg, inner_jpg):
    doc = Document()
    # cover section — no header bg
    s0 = doc.sections[0]
    s0.top_margin = Inches(0)
    s0.bottom_margin = Inches(0)
    s0.left_margin = Inches(0)
    s0.right_margin = Inches(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(cover_jpg), width=Inches(8.27), height=Inches(11.69))

    doc.add_section(WD_SECTION.NEW_PAGE)
    s1 = doc.sections[1]
    s1.top_margin = Inches(1.1)
    s1.bottom_margin = Inches(0.45)
    s1.left_margin = Inches(0.45)
    s1.right_margin = Inches(0.45)
    add_bg_to_header(s1, inner_jpg)

    MAX = 26
    for vi, (vertical, rows) in enumerate(groups):
        campaigns = [row_to_campaign(r) for r in rows]
        for i in range(0, len(campaigns), MAX):
            chunk = campaigns[i : i + MAX]
            add_table(doc, vertical if i == 0 else f"{vertical} (cont.)", chunk)
            if i + MAX < len(campaigns):
                doc.add_section(WD_SECTION.NEW_PAGE)
                sec = doc.sections[-1]
                sec.top_margin = Inches(1.1)
                add_bg_to_header(sec, inner_jpg)

    add_contact_block(doc)
    doc.save(OUT_DOCX)


def main():
    ASSETS.mkdir(exist_ok=True)
    cover = ASSETS / "cover-300dpi.jpg"
    inner = ASSETS / "inner-300dpi.jpg"

    print(f"Rendering artwork {PAGE_W}×{PAGE_H}px @ {DPI} DPI …")
    make_cover(cover)
    make_inner_bg(inner)

    groups = load_offers()
    total = sum(len(g[1]) for g in groups)
    print(f"{total} UK campaigns · {len(groups)} verticals")

    build_fresh_docx(groups, cover, inner)
    print("Built ConvertLane rate card (300 DPI artwork · dark-table layout)")

    shutil.copy(OUT_DOCX, OUT_DOWNLOADS)
    print(f"Saved → {OUT_DOCX}")
    print(f"Copied → {OUT_DOWNLOADS}")
    print(f"Cover JPEG: {cover.stat().st_size // 1024} KB · {DPI} DPI")


if __name__ == "__main__":
    main()
